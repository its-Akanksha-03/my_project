import os
import json
import fitz  # PyMuPDF
from docx import Document
import gradio as gr

# Load checklist
with open("checklist.json", "r") as f:
    checklist = json.load(f)

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def detect_doc_type(text):
    keywords = {
        "Articles of Association": "Articles of Association",
        "Memorandum of Association": "Memorandum of Association",
        "Board Resolution": "Board Resolution",
        "UBO Declaration": "UBO Declaration",
        "Register of Members and Directors": "Register of Members and Directors"
    }
    for key in keywords:
        if key.lower() in text.lower():
            return keywords[key]
    return "Unknown"

def load_reference_knowledge():
    knowledge = ""
    for filename in os.listdir("reference_docs"):
        path = os.path.join("reference_docs", filename)
        ext = os.path.splitext(filename)[1].lower()
        if ext == ".pdf":
            with fitz.open(path) as doc:
                for page in doc:
                    knowledge += page.get_text()
        elif ext == ".txt":
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                knowledge += f.read()
    return knowledge

def detect_red_flags(text):
    issues = []
    reference = load_reference_knowledge()

    if "UAE Federal Courts" in text:
        issues.append({
            "issue": "Incorrect jurisdiction reference",
            "suggestion": "Update to ADGM Courts",
            "severity": "High",
            "source": "ADGM Regulations 2020"
        })

    if "signatory" not in text.lower():
        issues.append({
            "issue": "Missing signatory section",
            "suggestion": "Add signatory section",
            "severity": "Medium",
            "source": "ADGM Document Templates"
        })

    if "limited liability" in text.lower() and "limited liability" not in reference.lower():
        issues.append({
            "issue": "Unverified clause: 'limited liability'",
            "suggestion": "Check against ADGM clause templates",
            "severity": "Low",
            "source": "Document Upload Categories"
        })

    return issues

def insert_comments(doc_path, issues):
    doc = Document(doc_path)
    if issues:
        para = doc.paragraphs[0]
        for issue in issues:
            comment = f"[Comment: {issue['issue']} — Suggestion: {issue['suggestion']} — Source: {issue.get('source', 'N/A')}]"
            para.text += f"\n{comment}"
    reviewed_path = "reviewed.docx"
    doc.save(reviewed_path)
    return reviewed_path

def review_docs(files):
    uploaded_types = []
    issues_found = []
    process = "Company Incorporation"
    required_docs = checklist[process]

    for file in files:
        ext = os.path.splitext(file.name)[1].lower()
        if ext != ".docx":
            continue
        text = extract_text_from_docx(file.name)
        doc_type = detect_doc_type(text)
        uploaded_types.append(doc_type)
        issues = detect_red_flags(text)
        if issues:
            issues_found.append({
                "document": doc_type,
                "issues": issues
            })
        insert_comments(file.name, issues)

    missing = [doc for doc in required_docs if doc not in uploaded_types]

    summary = {
        "process": process,
        "documents_uploaded": len(uploaded_types),
        "required_documents": len(required_docs),
        "missing_documents": missing,
        "issues_found": issues_found
    }

    with open("summary.json", "w") as f:
        json.dump(summary, f, indent=2)

    return f"Uploaded {len(uploaded_types)} documents.\nMissing: {', '.join(missing)}", "reviewed.docx", "summary.json"

gr.Interface(
    fn=review_docs,
    inputs=[
        gr.Files(label="Upload .docx Files", file_types=[".docx"])
    ],
    outputs=[
        gr.Textbox(label="Summary"),
        gr.File(label="Download Reviewed Document"),
        gr.File(label="Download JSON Summary")
    ],
    title="Corporate Agent – ADGM Legal Assistant"
).launch()
